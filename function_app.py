import re
import logging
import os
import tempfile
import time
import azure.functions as func
import redis
import numpy as np
from azure.storage.blob import BlobServiceClient
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_openai import OpenAIEmbeddings
from redis.commands.search.field import VectorField, TextField
from redis.commands.search.indexDefinition import IndexDefinition, IndexType
 
# Configurar logging settings
logging.basicConfig(level=logging.INFO)
 
# Função para sanitizar o nome do arquivo
def sanitize_filename(filename):
    # Substitui caracteres inválidos por "_"
    filename = re.sub(r'[^a-zA-Z0-9_\-\.]', '_', filename)
    # Remove pontos e underscores consecutivos
    filename = re.sub(r'\.+', '.', filename)
    filename = re.sub(r'_+', '_', filename)
    # Remove ponto ou underscore do início e fim do nome do arquivo
    filename = filename.strip('._ ')
    return filename
 
# Funções para extrair texto de diferentes formatos de arquivo
def extract_text_from_pdf(file_path):
    import fitz  # PyMuPDF
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text
 
def extract_text_from_docx(file_path):
    import docx
    doc = docx.Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text
 
def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()
 
# Função para gerar embeddings com tentativas de repetição em caso de falha
def get_embeddings_with_retry(texts, embeddings, max_retries=5, delay=2):
    retries = 0
    while retries < max_retries:
        try:
            return embeddings.embed_documents(texts)
        except Exception as e:
            if "Service Unavailable" in str(e):
                retries += 1
                time.sleep(delay)
            else:
                raise e
    raise Exception("Max retries exceeded for embedding service.")
 
# Função para criar o índice no Redis
def create_redis_index(redis_client, index_name, dimension):
    try:
        redis_client.ft(index_name).info()
    except:
        schema = [
            TextField("content"),
            VectorField("embedding", "FLAT", {"TYPE": "FLOAT32", "DIM": dimension, "DISTANCE_METRIC": "COSINE"})
        ]
        definition = IndexDefinition(prefix=["doc:"], index_type=IndexType.HASH)
        redis_client.ft(index_name).create_index(schema, definition=definition)
 
app = func.FunctionApp(http_auth_level=func.AuthLevel.FUNCTION)
 
 
@app.route(route="upload", methods=["POST"])
async def main(req: func.HttpRequest) -> func.HttpResponse:
    """
    Azure Function para receber arquivo, processar, gerar embeddings, criar índice no Redis e salvar no Blob Storage.
    """
    logging.info('Processing HTTP POST request')
 
    redis_client = None  # Inicializa a variável redis_client
 
    try:
        # Extrair parâmetros de conexão do Redis, OpenAI e configuração do processamento
        required_params = {
            'redis_host': req.form.get('redis_host'),
            'redis_port': req.form.get('redis_port'),
            'redis_password': req.form.get('redis_password'),
            'openai_embedding_key': req.form.get('openai_embedding_key'),
            'openai_embedding_model': req.form.get('openai_embedding_model'),
            'azure_storage_connection_string': req.form.get('azure_storage_connection_string'),
            'container_name': req.form.get('container_name'),
            'chunk_size': req.form.get('chunk_size'),
            'chunk_overlap': req.form.get('chunk_overlap'),
            'title': req.form.get('title'),
            'source': req.form.get('source')
        }
 
        # Verificar parâmetros faltantes
        missing_params = [key for key, value in required_params.items() if not value]
        if missing_params:
            logging.warning(f"Missing parameters: {', '.join(missing_params)}")
            return func.HttpResponse(f"Missing parameters: {', '.join(missing_params)}", status_code=400)
 
        # Converter redis_port, chunk_size e chunk_overlap para inteiros
        redis_port = int(required_params['redis_port'])
        chunk_size = int(required_params['chunk_size'])
        chunk_overlap = int(required_params['chunk_overlap'])
 
        # Configurar conexões e variáveis
        os.environ["OPENAI_API_KEY"] = required_params['openai_embedding_key']
        embeddings = OpenAIEmbeddings(model=required_params['openai_embedding_model'], openai_api_key=required_params['openai_embedding_key'])
 
        redis_client = redis.Redis(host=required_params['redis_host'], port=redis_port, password=required_params['redis_password'])
        blob_service_client = BlobServiceClient.from_connection_string(required_params['azure_storage_connection_string'])
        title = required_params['title']
        source = required_params['source']
 
        # Processar o arquivo enviado
        file = req.files.get('file')
        if not file:
            return func.HttpResponse("File is missing", status_code=400)
       
        file_name = sanitize_filename(file.filename)  # Sanitize the file name
        file_extension = os.path.splitext(file_name)[1]  # Extrai a extensão do arquivo
        temp_dir = tempfile.TemporaryDirectory()
        file_path = os.path.join(temp_dir.name, file_name)
 
        with open(file_path, 'wb') as f:
            f.write(file.stream.read())
 
        # Extração de conteúdo do arquivo
        if file_extension == '.pdf':
            content = extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            content = extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            content = extract_text_from_txt(file_path)
        else:
            return func.HttpResponse(f"Unsupported file type: {file_extension}", status_code=400)
 
        # Divisão do conteúdo em chunks e geração de embeddings
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        documents = [Document(page_content=chunk) for chunk in text_splitter.split_text(content)]
        texts = [doc.page_content for doc in documents]
        doc_embeddings = get_embeddings_with_retry(texts, embeddings)
 
        # Criar índice no Redis
        embedding_dimension = len(doc_embeddings[0])
        create_redis_index(redis_client, "document_index", embedding_dimension)
 
        # Armazenar os documentos e embeddings no Redis
        for i, doc_embedding in enumerate(doc_embeddings):
            doc_id = f"doc:{file_name}:{i}"
            embedding_array = np.array(doc_embedding, dtype=np.float32)
            redis_client.hset(doc_id, mapping={"content": ('Titulo: ' + title + '\n\n' + texts[i] + '\n\nFonte: '+ source), "embedding": embedding_array.tobytes()})
 
        # Upload do arquivo para o Blob Storage
        container_name = required_params['container_name']
        blob_client = blob_service_client.get_blob_client(container=container_name, blob=file_name)
        with open(file_path, "rb") as data:
            blob_client.upload_blob(data, overwrite=True)
 
        # Fechar o diretório temporário
        temp_dir.cleanup()
 
        logging.info("File processed, indexed, and uploaded successfully.")
        return func.HttpResponse(f"{file_name}", status_code=200)
 
    except Exception as e:
        logging.error(f"Failed to process file: {e}")
        return func.HttpResponse(f"An error occurred: {str(e)}", status_code=500)
 
    finally:
        if redis_client:
            redis_client.close()
            logging.info("Redis connection closed")
 
@app.route(route="list-files", methods=["GET"])
async def list_files(req: func.HttpRequest) -> func.HttpResponse:
    required_params = {
            'azure_storage_connection_string': req.form.get('azure_storage_connection_string'),
            'container_name': req.form.get('container_name')
        }
 
   
    missing_params = [key for key, value in required_params.items() if not value]
    if missing_params:
        logging.warning(f"Missing parameters: {', '.join(missing_params)}")
        return func.HttpResponse(f"Missing parameters: {', '.join(missing_params)}", status_code=400)
 
    try:
        blob_service_client = BlobServiceClient.from_connection_string(required_params['azure_storage_connection_string'])
        container_client = blob_service_client.get_container_client(required_params['container_name'])
        blob_list = container_client.list_blobs()
 
        files = [blob.name for blob in blob_list]
        return func.HttpResponse(f"Files: {', '.join(files)}", status_code=200)
    except Exception as e:
        logging.error(f"Error listing files: {e}")
        return func.HttpResponse(f"An error occurred while listing files: {str(e)}", status_code=500)
 
 
@app.route(route="delete-file", methods=["DELETE"])
async def delete_file(req: func.HttpRequest) -> func.HttpResponse:
    required_params = {
            'redis_host': req.form.get('redis_host'),
            'redis_port': req.form.get('redis_port'),
            'redis_password': req.form.get('redis_password'),
            'azure_storage_connection_string': req.form.get('azure_storage_connection_string'),
            'container_name': req.form.get('container_name'),
            'file_name': req.form.get('file_name')
        }
 
   
    missing_params = [key for key, value in required_params.items() if not value]
    if missing_params:
        logging.warning(f"Missing parameters: {', '.join(missing_params)}")
        return func.HttpResponse(f"Missing parameters: {', '.join(missing_params)}", status_code=400)
 
    try:
        # Conexão ao Redis
        redis_client = redis.Redis(host=required_params['redis_host'], port=int(required_params['redis_port']), password=required_params['redis_password'])
        # Deletar documentos do Redis
        redis_keys = redis_client.keys(f"doc:{sanitize_filename(required_params['file_name'])}:*")
        for key in redis_keys:
            redis_client.delete(key)
 
        # Deletar arquivo do Blob Storage
        blob_service_client = BlobServiceClient.from_connection_string(required_params['azure_storage_connection_string'])
        blob_client = blob_service_client.get_blob_client(container=required_params['container_name'], blob=required_params['file_name'])
        blob_client.delete_blob()
 
        return func.HttpResponse(f"File {required_params['file_name']} and associated Redis entries deleted successfully", status_code=200)
    except Exception as e:
        logging.error(f"Failed to delete file: {e}")
        return func.HttpResponse(f"An error occurred: {str(e)}", status_code=500)
    finally:
        if redis_client:
            redis_client.close()
            logging.info("Redis connection closed")